# app.py - API Flask: PDF -> XLSX / DOCX
# /convert?to=excel|word&mode=auto|rpt|table|form&engine=auto|plumber|camelot|tabula|ocr

import io, os, re, tempfile, warnings, statistics, importlib, datetime as _dt
from typing import List, Tuple, Optional
from difflib import SequenceMatcher

from flask import Flask, request, send_file, abort, make_response
from flask_cors import CORS
import pdfplumber
from docx import Document
from openpyxl import Workbook

warnings.filterwarnings("ignore", category=DeprecationWarning)
app = Flask(__name__)

# ---------------- CONFIG ----------------
MAX_MB = 20
ALLOWED_ORIGINS = {"https://natanjs01.github.io"}  # ajuste se precisar

CORS(
    app,
    resources={r"/convert": {"origins": list(ALLOWED_ORIGINS)}},
    methods=["POST", "OPTIONS"],
    allow_headers=["Content-Type"],
    max_age=86400,
)

@app.after_request
def add_cors_headers(resp):
    origin = request.headers.get("Origin", "")
    if origin in ALLOWED_ORIGINS:
        resp.headers["Access-Control-Allow-Origin"] = origin
        resp.headers["Vary"] = "Origin"
        resp.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return resp

# endpoint leve para health/keep-alive
@app.get("/health")
def health():
    return "ok", 200

# ---------------- WORD ----------------
def pdf_to_docx(file_stream) -> bytes:
    doc = Document()
    with pdfplumber.open(file_stream) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if i > 1:
                doc.add_page_break()
            doc.add_heading(f"Página {i}", level=2)
            for line in text.splitlines():
                doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.read()

# ---------------- EXCEL CORE ----------------
EXCEL_MAX_ROWS = 1_048_576
EXCEL_MAX_COLS = 16_384
CELL_MAX = 32_000  # (< 32767 seguro)
_xml_illegal_re = re.compile(u"[\u0000-\u0008\u000b\u000c\u000e-\u001f\uD800-\uDFFF\uFFFE\uFFFF]")

def _sanitize_cell(s: str) -> str:
    if s is None: return ""
    if not isinstance(s, str): s = str(s)
    s = _xml_illegal_re.sub("", s).replace("\xa0", " ")
    s = " ".join(s.split())
    return s[:CELL_MAX]

def _ensure_sheet_capacity(wb, ws_name_base, ws, add_rows, add_cols):
    r, c = ws.max_row or 0, ws.max_column or 0
    need = (r + add_rows > EXCEL_MAX_ROWS) or (max(c, add_cols) > EXCEL_MAX_COLS)
    if not need: return ws
    i = 2
    names = {s.title for s in wb.worksheets}
    while f"{ws_name_base}_{i}" in names: i += 1
    return wb.create_sheet(f"{ws_name_base}_{i}")

def _append_rows(wb, ws, ws_base, rows):
    for row in rows:
        row = [_sanitize_cell(x) if not isinstance(x, (_dt.date, int, float)) else x for x in row]
        ws2 = _ensure_sheet_capacity(wb, ws_base, ws, 1, len(row))
        if ws2 is not ws: ws = ws2
        ws.append(row)
    return ws

def _words(page):
    w = page.extract_words(keep_blank_chars=False, use_text_flow=True) or []
    return [it for it in w if _sanitize_cell(it.get("text"))]

def _group_by_y(words, y_tol=3):
    rows = {}
    for w in words: rows.setdefault(round(w["top"]/y_tol), []).append(w)
    out = []
    for _, items in sorted(rows.items(), key=lambda kv: kv[0]):
        items.sort(key=lambda w: w["x0"]); out.append(items)
    return out

def _page_text(page) -> str:
    return (page.extract_text() or "").replace("\xa0", " ")

# -------- Detectores de modelos --------
def _detect_rpt_lojas(text: str) -> bool:
    return ("LOJA" in text and "CHAPA" in text and "FUNÇÃO" in text and "VALOR" in text)
def _detect_rpt_desligados(text: str) -> bool:
    return ("Relatório de Colaboradores" in text and "Desligados" in text and "Nome" in text and "Cpf" in text)

# -------- Geração de grade --------
def _build_grid_by_gaps(words, max_cols=20):
    if not words: return [0,1e9], lambda x:0
    xs = sorted(w["x0"] for w in words)
    gaps = [xs[i+1]-xs[i] for i in range(len(xs)-1)]
    if not gaps: return [0,1e9], lambda x:0
    med = statistics.median(gaps)
    p90 = sorted(gaps)[int(len(gaps)*0.90)] if len(gaps) >= 10 else max(gaps)
    thr = max(med*2.6, p90)
    bounds = [xs[0]-8]
    for i,g in enumerate(gaps):
        if g >= thr: bounds.append(xs[i] + g/2)
    bounds.append(xs[-1]+8)
    if len(bounds)-1 > max_cols:
        step=(len(bounds)-1)/max_cols; nb=[bounds[0]]; acc=0
        for _ in range(max_cols-1):
            acc+=step; nb.append(bounds[int(round(acc))])
        nb.append(bounds[-1]); bounds=nb
    def cidx(x):
        lo,hi=0,len(bounds)-1
        while lo<hi:
            m=(lo+hi)//2
            if x<bounds[m]: hi=m
            else: lo=m+1
        return max(0,min(lo-1,len(bounds)-2))
    return bounds,cidx

def _build_grid_from_header(words, header_tokens: List[str]):
    def norm(s): return _sanitize_cell(s).lower()
    header_pos={}
    for w in words:
        t=norm(w["text"])
        for token in header_tokens:
            if token in t and token not in header_pos:
                header_pos[token]=(w["x0"]+w["x1"]) / 2.0
    if len(header_pos)>=2:
        xs=sorted(header_pos[t] for t in header_tokens if t in header_pos)
        bounds=[xs[0]-30.0]
        for i in range(len(xs)-1): bounds.append((xs[i]+xs[i+1])/2.0)
        bounds.append(xs[-1]+80.0)
        def col_index(x):
            lo,hi=0,len(bounds)-1
            while lo<hi:
                mid=(lo+hi)//2
                if x<bounds[mid]: hi=mid
                else: lo=mid+1
            return max(0,min(lo-1,len(bounds)-2))
        return bounds,col_index
    return _build_grid_by_gaps(words)

# -------- Materializadores específicos --------
_HEADERS_LOJAS = ["chapa","nome","funç","ref","valor","sind"]

def _extract_loja_name(line_words: List[dict]) -> str:
    txt = " ".join(_sanitize_cell(w["text"]) for w in line_words)
    m = re.search(r"LOJA\s+\d+\s*=\s*(.+)$", txt, flags=re.I)
    return m.group(1).strip() if m else txt.strip()

def _materialize_rpt_lojas(words) -> List[List[str]]:
    lines=_group_by_y(words); loja_atual=""; rows=[]
    bounds,cidx=_build_grid_from_header(words,_HEADERS_LOJAS)
    for line in lines:
        txt=" ".join(_sanitize_cell(w["text"]) for w in line); up=txt.upper()
        if up.startswith("LOJA "): loja_atual=_extract_loja_name(line); continue
        if "TOTAL" in up or up.startswith("PÁGINA") or up.startswith("RELAT") or up.startswith("DATA:"): continue
        ncols=max(1,len(bounds)-1); cells=[""]*ncols
        for w in line:
            col=cidx(w["x0"]); cells[col]=(cells[col]+" "+_sanitize_cell(w["text"])).strip()
        if any(cells):
            while len(cells)<6: cells.append("")
            rows.append([loja_atual]+cells[:6])
    if rows: rows.insert(0,["Loja","Chapa","Nome","Função","Ref","Valor","Sind"])
    return rows

_HEADERS_DESL = ["nome","cpf","admiss","demiss","filial","chapa"]

def _materialize_rpt_desligados(words) -> List[List[str]]:
    bounds,cidx=_build_grid_from_header(words,_HEADERS_DESL)
    lines=_group_by_y(words); rows=[]
    for line in lines:
        txt=" ".join(_sanitize_cell(w["text"]) for w in line)
        if "Relatório de Colaboradores" in txt or "CNPJ:" in txt or "PAG.:" in txt: continue
        if txt.startswith("Página") or txt.startswith("Rel:"): continue
        ncols=max(1,len(bounds)-1); cells=[""]*ncols
        for w in line:
            col=cidx(w["x0"]); cells[col]=(cells[col]+" "+_sanitize_cell(w["text"])).strip()
        if any(cells):
            while len(cells)<6: cells.append("")
            rows.append(cells[:6])
    if rows: rows.insert(0,["Nome","CPF","Dt.Admissão","Dt.Demissão","Filial","Chapa"])
    return rows

# -------- Genéricos --------
_LABEL_HINTS = {"nome","placa","data","telefone","modelo","montadora","ano","km","código","descricao","descrição","abs","airbag","injeção"}

def _is_probably_label(text):
    t=_sanitize_cell(text).lower().rstrip(":")
    return (text.endswith(":")) or (t in _LABEL_HINTS and len(text)<=25)

def _materialize_form(words):
    lines=_group_by_y(words); rows=[]
    for line in lines:
        i=0
        while i<len(line):
            t=_sanitize_cell(line[i]["text"])
            if _is_probably_label(t):
                campo=t.rstrip(":"); j=i+1; vals=[]
                while j<len(line):
                    tj=_sanitize_cell(line[j]["text"])
                    if _is_probably_label(tj): break
                    vals.append(tj); j+=1
                rows.append([campo," ".join(vals).strip()]); i=j
            else:
                i+=1
    if not rows:
        for line in lines:
            rows.append([" ".join(_sanitize_cell(w["text"]) for w in line)])
    return rows

def _build_grid_by_gaps_table(words):
    bounds,cidx=_build_grid_by_gaps(words)
    lines=_group_by_y(words); table=[]; ncols=max(1,len(bounds)-1)
    for line in lines:
        cells=[""]*ncols; prev_col=prev_x1=None
        for w in line:
            col=cidx(w["x0"]); text=_sanitize_cell(w["text"])
            if prev_col==col and prev_x1 is not None:
                cells[col]=(cells[col]+" "+text).strip(); prev_x1=w["x1"]
            else:
                cells[col]=(cells[col]+" "+text).strip() if cells[col] else text
                prev_col,prev_x1=col,w["x1"]
        if any(cells): table.append(cells)
    return table

def _auto_mode(words, text):
    if not words: return "empty"
    colon=text.count(":"); uniq_x=len({int(w["x0"]//10) for w in words})
    lines=_group_by_y(words); avg_len=sum(len(l) for l in lines)/max(1,len(lines))
    if colon>=8 and uniq_x>=20: return "form"
    if avg_len>=6: return "table"
    return "form" if colon>=3 else "table"

# -------- Tipagem, formatação e utilidades --------
_PT_MONTHS = {"jan":1,"fev":2,"mar":3,"abr":4,"mai":5,"jun":6,"jul":7,"ago":8,"set":9,"out":10,"nov":11,"dez":12}

def _to_number_pt(s: str):
    t = _sanitize_cell(s)
    if not t: return None
    t2 = t.replace(".", "")
    if "," in t2:
        t2 = t2.replace(",", ".")
    try:
        v = float(t2)
        if abs(v - round(v)) < 1e-9:
            return int(round(v))
        return v
    except Exception:
        return None

def _to_date_pt(s: str):
    t = _sanitize_cell(s)
    if not t: return None
    t = t.replace("-", "/").replace(".", "/")
    m = re.match(r"^\s*(\d{1,2})/(\d{1,2})/(\d{2,4})\s*$", t)
    if m:
        d, mo, y = map(int, m.groups())
        if y < 100: y += 2000 if y < 70 else 1900
        try: return _dt.date(y, mo, d)
        except: return None
    m = re.match(r"^\s*(\d{1,2})/([A-Za-z]{3})/(\d{2,4})\s*$", t)
    if m:
        d, mon_txt, y = m.groups()
        d = int(d); y = int(y); 
        if y < 100: y += 2000 if y < 70 else 1900
        mon = _PT_MONTHS.get(mon_txt.lower())
        if mon:
            try: return _dt.date(y, mon, d)
            except: return None
    return None

def _guess_types_in_row(row: list):
    out=[]
    for c in row:
        if c is None or c == "": out.append(""); continue
        d = _to_date_pt(c)
        if d: out.append(d); continue
        n = _to_number_pt(c)
        out.append(n if n is not None else _sanitize_cell(c))
    return out

def _autowidth(ws):
    from openpyxl.utils import get_column_letter
    widths={}
    for r in ws.iter_rows(values_only=True):
        for i,val in enumerate(r, start=1):
            s = str(val) if val is not None else ""
            w = min(60, max(8, int(len(s)*0.9)+2))
            widths[i] = max(widths.get(i, 10), w)
    for i,w in widths.items():
        ws.column_dimensions[get_column_letter(i)].width = w

def _style_header(ws, header_row=1):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    thin = Side(style="thin", color="DDDDDD")
    for cell in ws[header_row]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="E8EEF9")
        cell.alignment = Alignment(vertical="center")
        cell.border = Border(top=thin,bottom=thin,left=thin,right=thin)
    ws.freeze_panes = f"A{header_row+1}"

def _fuzzy_ratio(a: list, b: list) -> float:
    sa = " | ".join(_sanitize_cell(x).lower() for x in a)
    sb = " | ".join(_sanitize_cell(x).lower() for x in b)
    return SequenceMatcher(None, sa, sb).ratio()

def _same_header(a: list, b: list, tol=0.78) -> bool:
    if not a or not b: return False
    return _fuzzy_ratio(a, b) >= tol

def _is_page_noise(row: list) -> bool:
    s = " ".join(_sanitize_cell(c).lower() for c in row)
    return bool(re.search(r"^p[aá]gina\b|\brel:|\bcnpj:|^data:\b|^hora:\b", s))

# -------- Cabeçalho inteligente / Normalizadores --------
def _looks_like_header(row, tokens: set) -> bool:
    if not row: return False
    s=" ".join(_sanitize_cell(c).lower() for c in row)
    hits=sum(1 for t in tokens if t in s)
    return hits >= max(3, (len(tokens)+1)//2)

def _inject_header_if_missing(rows, expected_header, tokens: set):
    if not rows: return rows
    i0=0
    while i0<len(rows) and not any(_sanitize_cell(c) for c in rows[i0]): i0+=1
    if i0>=len(rows): return rows
    if _looks_like_header(rows[i0], tokens): return rows
    return rows[:i0] + [expected_header] + rows[i0:]

_LOJA_HEADER = ["Loja","Chapa","Nome","Função","Ref","Valor","Sind"]
_LOJA_TOKENS = {"loja","chapa","nome","funç","ref","valor","sind"}

def _normalize_rpt_lojas_rows(rows: List[List[str]], prev_loja: str | None):
    """
    Remove 'LOJA n = ...' e linhas vazias, injeta cabeçalho se faltar,
    garante 7 colunas e preenche 'Loja' usando o último marcador visto.
    """
    cur_loja = prev_loja or ""
    out: List[List[str]] = []

    for r in rows:
        joined = " ".join(_sanitize_cell(c) for c in r if c).strip()
        if not joined:
            continue
        # marcador de loja?
        m = re.search(r"loja\s+\d+\s*=\s*(.+)", joined, flags=re.I)
        if m:
            cur_loja = m.group(1).strip()
            continue
        rr = list(r)
        while len(rr) < 7:
            rr.append("")
        if not _sanitize_cell(rr[0]):
            rr[0] = cur_loja or prev_loja or ""
        out.append(rr[:7])

    out = _inject_header_if_missing(out, _LOJA_HEADER, _LOJA_TOKENS)
    return out, (cur_loja or prev_loja or "")

# -------- “Dupla coluna” para plumber --------
def _plumber_two_columns(words_all, text, mode_hint):
    if not words_all: return []
    xs = sorted((w["x0"] + w["x1"]) / 2 for w in words_all)
    if len(xs) < 40:
        split_x = (min(xs)+max(xs))/2
    else:
        p40 = xs[int(0.40*len(xs))]; p60 = xs[int(0.60*len(xs))]
        split_x = (p40 + p60) / 2.0
    left  = [w for w in words_all if ((w["x0"] + w["x1"])/2) <= split_x]
    right = [w for w in words_all if ((w["x0"] + w["x1"])/2) >  split_x]

    def extract_from(words_local):
        if mode_hint == "form":   return _materialize_form(words_local)
        if mode_hint == "table":  return _build_grid_by_gaps_table(words_local)
        if _detect_rpt_lojas(text):      return _materialize_rpt_lojas(words_local)
        if _detect_rpt_desligados(text): return _materialize_rpt_desligados(words_local)
        mm = _auto_mode(words_local, text)
        return _materialize_form(words_local) if mm=="form" else _build_grid_by_gaps_table(words_local)

    if len(left) > 50 and len(right) > 50:
        rows_left  = extract_from(left) or []
        rows_right = extract_from(right) or []
        out = []
        if rows_left:  out.extend(rows_left)
        if rows_left and rows_right: out.append([])
        if rows_right: out.extend(rows_right)
        return out
    return extract_from(words_all)

# ---------------- Multi-engine ----------------
def _lib_available(modname: str) -> bool:
    try: importlib.import_module(modname); return True
    except Exception: return False

def _score_table(rows: List[List[str]]) -> float:
    if not rows: return 0.0
    n=len(rows); m=max(len(r) for r in rows)
    if n==0 or m==0: return 0.0
    total=n*m; filled=sum(1 for r in rows for c in r if _sanitize_cell(c))
    density=filled/max(1,total)
    header_boost=0.0; hdr=" ".join(_sanitize_cell(x).lower() for x in rows[0])
    for k in ("chapa","nome","cpf","filial","funç","ref","valor","sind","data","descri"):
        if k in hdr: header_boost += 0.05
    return (n**0.7)*(m**0.6)*(0.5+0.5*density)*(1.0+header_boost)

def _engine_plumber(page, text, mode_hint="auto") -> List[List[str]]:
    words_all=_words(page)
    if not words_all: return []
    return _plumber_two_columns(words_all, text, mode_hint)

def _engine_camelot(pdf_path: str, page_number: int) -> List[List[List[str]]]:
    """Retorna lista de TABELAS (cada tabela é List[List[str]])."""
    if not _lib_available("camelot"): return []
    import camelot
    p=str(page_number); out=[]
    for flavor in ("lattice","stream"):
        try:
            tables=camelot.read_pdf(pdf_path, pages=p, flavor=flavor, strip_text="\n")
        except Exception:
            continue
        for t in getattr(tables,"tables",[]):
            rows=[list(map(_sanitize_cell,row)) for row in t.df.values.tolist()]
            if len(rows) >= 2 and max(len(r) for r in rows) >= 2:
                out.append(rows)
    return out

def _engine_tabula(pdf_path: str, page_number: int) -> List[List[List[str]]]:
    """Retorna lista de TABELAS (cada tabela é List[List[str]])."""
    if not _lib_available("tabula"): return []
    import tabula
    p=str(page_number); out=[]
    for lattice in (True, False):
        try:
            dfs=tabula.read_pdf(pdf_path, pages=p, lattice=lattice, stream=(not lattice), multiple_tables=True)
        except Exception:
            continue
        for df in dfs or []:
            rows=[list(map(_sanitize_cell, map(str,row))) for row in df.values.tolist()]
            if len(rows) >= 2 and max(len(r) for r in rows) >= 2:
                out.append(rows)
    return out

def _engine_ocr(pdf_path: str, page_number: int) -> List[List[str]]:
    if not (_lib_available("pdf2image") and _lib_available("pytesseract")): return []
    from pdf2image import convert_from_path
    import pytesseract
    try:
        imgs=convert_from_path(pdf_path, first_page=page_number, last_page=page_number, dpi=300)
        if not imgs: return []
        img=imgs[0]; data=pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, lang="por+eng")
    except Exception: return []
    rows=[]; current=[]; last_line=None
    for i in range(len(data["text"])):
        if int(data["conf"][i]) < 0: continue
        text=_sanitize_cell(data["text"][i])
        if not text: continue
        line=data["line_num"][i]; x=data["left"][i]
        if last_line is None: last_line=line
        if line != last_line:
            if current:
                current.sort(key=lambda z: z[0]); rows.append([t for _,t in current])
            current=[]; last_line=line
        current.append((x,text))
    if current:
        current.sort(key=lambda z: z[0]); rows.append([t for _,t in current])
    return rows

def _is_scanned_page(page) -> bool:
    try: num_imgs=len(page.images or [])
    except Exception: num_imgs=0
    words=_words(page)
    return num_imgs>=1 and len(words)<10

def _extract_best_for_page(pdf_path: str, page, text: str, mode_hint="auto", engine="auto") -> List[List[str]]:
    if engine == "plumber":
        cands=[("plumber", lambda: _engine_plumber(page,text,mode_hint))]
    elif engine == "camelot":
        cands=[("camelot", lambda: _engine_camelot(pdf_path,page.page_number))]
    elif engine == "tabula":
        cands=[("tabula", lambda: _engine_tabula(pdf_path,page.page_number))]
    elif engine == "ocr":
        cands=[("ocr", lambda: _engine_ocr(pdf_path,page.page_number))]
    else:
        if _is_scanned_page(page):
            cands=[("ocr", lambda: _engine_ocr(pdf_path,page.page_number))]
        else:
            cands=[
                ("plumber", lambda: _engine_plumber(page,text,mode_hint)),
                ("camelot", lambda: _engine_camelot(pdf_path,page.page_number)),
                ("tabula",  lambda: _engine_tabula(pdf_path,page.page_number)),
            ]
    best_rows=[]; best=0.0
    for name, fn in cands:
        try:
            cand = fn() or []
            if name in ("camelot","tabula"):
                page_rows=[]
                for i, tbl in enumerate(cand):
                    if i>0: page_rows.append([])
                    page_rows.extend(tbl)
                rows = page_rows
            else:
                rows = cand
            s = _score_table(rows)
            if s > best:
                best, best_rows = s, rows
        except Exception:
            continue
    return best_rows

# ---------------- Assinatura/Fuzzy header ----------------
def _sig_from_header(row: List[str]) -> Tuple[int, Tuple[str, ...]]:
    norm = tuple(_sanitize_cell(c).lower() for c in row)
    return (len(row), norm)

def _is_header_row(row: List[str], known_tokens: Optional[set]=None) -> bool:
    s=" ".join(_sanitize_cell(c).lower() for c in row)
    toks = known_tokens or {"loja","chapa","nome","funç","ref","valor","sind","cpf","filial","admiss","demiss","data","descri","total"}
    hits = sum(1 for t in toks if t in s)
    return hits >= 2

# ---------------- PDF -> EXCEL (agregador contínuo) ----------------
def pdf_to_excel(file_stream_or_path, force_mode: str = "auto", engine: str = "auto") -> bytes:
    """
    force_mode: 'auto' | 'rpt' | 'table' | 'form'
    engine:     'auto' | 'plumber' | 'camelot' | 'tabula' | 'ocr'
    """
    from openpyxl.styles import numbers

    wb = Workbook(); ws = wb.active; ws.title = "Dados"; base = "Dados"

    need_cleanup=False
    if isinstance(file_stream_or_path,(str,os.PathLike)):
        pdf_path=str(file_stream_or_path)
    else:
        fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
        with os.fdopen(fd,"wb") as tmp: tmp.write(file_stream_or_path.read())
        file_stream_or_path.seek(0); need_cleanup=True

    loja_state = ""                       # RPT lojas
    current_header: List[str] = []        # último cabeçalho “vigente”
    header_written = False                # já escrevi cabeçalho no Excel?
    header_row_index = None               # linha do cabeçalho no Excel
    known_tokens = {"loja","chapa","nome","funç","ref","valor","sind","cpf","filial","admiss","demiss","data","descri","total"}

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text=_page_text(page)
                mode_hint = force_mode if force_mode in {"form","table"} else "auto"

                rows = _extract_best_for_page(pdf_path, page, text, mode_hint, engine)
                if not rows:
                    continue

                # 1) filtros de ruído (remover headers de página)
                rows = [r for r in rows if not _is_page_noise(r)]

                # 2) normalizações específicas
                if _detect_rpt_lojas(text):
                    rows, loja_state = _normalize_rpt_lojas_rows(rows, loja_state)
                elif _detect_rpt_desligados(text):
                    expected=["Nome","CPF","Dt.Admissão","Dt.Demissão","Filial","Chapa"]
                    tokens={"nome","cpf","admiss","demiss","filial","chapa"}
                    rows=_inject_header_if_missing(rows, expected, tokens)

                # 3) detectar cabeçalho
                i = 0
                while i < len(rows) and not any(_sanitize_cell(c) for c in rows[i]): i += 1
                if i < len(rows) and _is_header_row(rows[i], known_tokens):
                    header = [ _sanitize_cell(c) for c in rows[i] ]
                    if not header_written:
                        ws = _append_rows(wb, ws, base, [header]); header_written=True
                        header_row_index = ws.max_row
                        current_header = header
                    else:
                        if not _same_header(header, current_header):
                            ws = _append_rows(wb, ws, base, [[]])
                            ws = _append_rows(wb, ws, base, [header])
                            header_row_index = ws.max_row
                            current_header = header
                    i += 1

                # 4) dados alinhados ao cabeçalho atual + tipagem
                data_rows = [r for r in rows[i:] if any(_sanitize_cell(c) for c in r)]
                for r in data_rows:
                    rr = _guess_types_in_row(r)
                    if current_header:
                        cols = len(current_header)
                        if len(rr) < cols: rr = rr + [""]*(cols-len(rr))
                        elif len(rr) > cols: rr = rr[:cols]
                    ws = _append_rows(wb, ws, base, [rr])

    finally:
        if need_cleanup and os.path.exists(pdf_path): os.remove(pdf_path)

    # 5) pós-formatação (estilo cabeçalho, formato numérico/data, largura)
    if header_written and header_row_index:
        _style_header(ws, header_row_index)
        for row in ws.iter_rows(min_row=header_row_index+1, values_only=False):
            for cell in row:
                if isinstance(cell.value, _dt.date):
                    cell.number_format = numbers.FORMAT_DATE_DDMMYYYY
                elif isinstance(cell.value, (int, float)):
                    cell.number_format = '#,##0.00'
    _autowidth(ws)

    bio=io.BytesIO(); wb.save(bio); bio.seek(0)
    return bio.read()

# ---------------- ROUTES ----------------
@app.route("/convert", methods=["OPTIONS"])
def convert_options(): return make_response("",204)

@app.post("/convert")
def convert():
    to_fmt = (request.args.get("to") or "").lower()
    mode   = (request.args.get("mode") or "auto").lower()
    engine = (request.args.get("engine") or "auto").lower()
    f = request.files.get("file")

    if to_fmt not in {"excel","word"}: return abort(400, "Parâmetro 'to' precisa ser 'excel' ou 'word'.")
    if not f: return abort(400, "Envie o arquivo no campo 'file'.")
    if f.mimetype != "application/pdf": return abort(400, "Envie um PDF válido.")
    f.seek(0,os.SEEK_END); size=f.tell(); f.seek(0)
    if size > MAX_MB*1024*1024: return abort(413, f"Arquivo excede {MAX_MB} MB.")

    try:
        if to_fmt=="word":
            out=pdf_to_docx(f.stream)
            return send_file(io.BytesIO(out),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=(f.filename or "arquivo").rsplit(".",1)[0]+".docx")
        else:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(f.read()); tmp_path=tmp.name
            try:
                with open(tmp_path,"rb") as fp:
                    out=pdf_to_excel(fp, force_mode=mode, engine=engine)
            finally:
                if os.path.exists(tmp_path): os.remove(tmp_path)
            return send_file(io.BytesIO(out),
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                as_attachment=True,
                download_name=(f.filename or "arquivo").rsplit(".",1)[0]+".xlsx")
    except Exception as e:
        return abort(500, f"Falha na conversão: {e}")

# ---------------- MAIN ----------------
if __name__ == "__main__":
    port=int(os.environ.get("PORT",8000))
    app.run(host="0.0.0.0", port=port)
